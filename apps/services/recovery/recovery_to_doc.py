
import os
from copy import deepcopy

from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

from table_process import HtmlToDocx



def convert_info_docx(img, res, save_folder, img_name):
    folder = os.path.join(save_folder,img_name)
    os.makedirs(folder,exist_ok=True)
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)

    flag = 1
    for i, region in enumerate(res):
        entity_type = region['type'].lower()
        if entity_type in ["header","footer"]:
            continue
        # if len(region['res']) == 0:
        #     continue
        img_idx = region['img_idx']
        if flag == 2 and region['layout'] == 'single':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
            flag = 1
        elif flag == 1 and region['layout'] == 'double':
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
            flag = 2

        if region['type'].lower() == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder,
                                    '{}_{}.jpg'.format(region['bbox'], img_idx))
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
        elif region['type'].lower() == 'title':
            doc.add_heading(region['res'][0]['text'])
        elif region['type'].lower() == 'table':
            parser = HtmlToDocx()
            parser.table_style = 'TableGrid'
            parser.handle_table(region['res']['html'], doc)
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            for i, line in enumerate(region['res']):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                text_run = paragraph.add_run(line['text'] + ' ')
                text_run.font.size = shared.Pt(10)

    # save to docx
    docx_path = os.path.join(save_folder,img_name, '{}.docx'.format(img_name))
    
    doc.save(docx_path)
    print('docx save to {}'.format(docx_path))

def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]['layout'] = 'single'
        return res

    sorted_boxes = sorted(res, key=lambda x: (x['bbox'][1], x['bbox'][0]))
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if _boxes[i]['bbox'][1] > _boxes[i - 1]['bbox'][3] and _boxes[i][
                    'bbox'][0] < w / 2 and _boxes[i]['bbox'][2] > w / 2:
                new_res += res_left
                new_res += res_right
                _boxes[i]['layout'] = 'single'
                new_res.append(_boxes[i])
            else:
                if _boxes[i]['bbox'][2] > w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]['bbox'][0] < w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]['bbox'][0] < w / 4 and _boxes[i]['bbox'][2] < 3 * w / 4:
            _boxes[i]['layout'] = 'double'
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]['bbox'][0] > w / 4 and _boxes[i]['bbox'][2] > w / 2:
            _boxes[i]['layout'] = 'double'
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]['layout'] = 'single'
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res

def merge_text_in_line(text_list,threshold = 4):
    if len(text_list) <= 1:
        return text_list
    sorted(text_list,key=lambda x:(x['text_region'][0][1],x['text_region'][0][0]))
    new_text_list = []
    idx = 0 
    while idx < len(text_list):
        idx_copy = idx
        _tmp_text = deepcopy(text_list[idx])
        for jdx in range(idx+1,len(text_list)):
            xmin = _tmp_text["text_region"][0][0]
            ymin = _tmp_text["text_region"][0][1]
            xmax = _tmp_text["text_region"][2][0]
            ymax = _tmp_text["text_region"][2][1]
            _xmin = text_list[jdx]["text_region"][0][0]
            _ymin = text_list[jdx]["text_region"][0][1]
            _xmax = text_list[jdx]["text_region"][2][0]
            _ymax = text_list[jdx]["text_region"][2][1]
            idx = jdx
            if abs(ymin-_ymin) <= threshold and abs(ymax-_ymax) <= threshold:
                # merge
                new_region = [
                    [min(xmin,_xmin),min(ymin,_ymin)],
                    [max(xmax,_xmax),min(ymin,_ymin)],
                    [max(xmax,_xmax),max(ymax,_ymax)],
                    [min(xmin,_xmin),max(ymax,_ymax)]
                ]
                new_text = _tmp_text["text"] + text_list[jdx]["text"]
                _tmp_text["text"]=new_text
                _tmp_text["text_region"]=new_region
                # print(f"_tmp_text:{_tmp_text}")
                idx = jdx
            else:
                break
        if idx == idx_copy:
            idx += 1
        new_text_list.append(_tmp_text) 
        # if idx == len(text_list)-1:
        #     break
    return new_text_list


if __name__ == "__main__":
    # data = [
    #     {'text_region': [[35.0, 9.0], [405.0, 10.0], [405.0, 24.0], [34.0, 23.0]], 'text': '从表1可看到刊载我院论文的期刊较多集中在湖北省武汉市地', 'confidence': 0.9976343512535095}, 
    #     {'text_region': [[8.0, 31.0], [405.0, 31.0], [405.0, 43.0], [8.0, 43.0]], 'text': '区，主要包括湖北中医杂志、中西医结合肝病杂志、中国中医骨伤科', 'confidence': 0.9853870868682861}, 
    #     {'text_region': [[8.0, 50.0], [405.0, 50.0], [405.0, 63.0], [8.0, 63.0]], 'text': '杂志等核心期刊，而湖北中医药大学学报、家庭保健、中西医结合研', 'confidence': 0.9875466227531433}, 
    #     {'text_region': [[7.0, 70.0], [406.0, 70.0], [406.0, 83.0], [7.0, 83.0]], 'text': '究、中国临床护理（又名：中国医学文摘：护理学）等为非核心期刊。', 'confidence': 0.9875941276550293}, 
    #     {'text_region': [[8.0, 90.0], [405.0, 90.0], [405.0, 103.0], [8.0, 103.0]], 'text': '从图2则可看到我院刊载在核心期刊与非核心期刊的发文量占有', 'confidence': 0.9912272095680237},
    #     {'text_region': [[9.0, 109.0], [176.0, 109.0], [176.0, 122.0], [9.0, 122.0]], 'text': '比例近三年没有显著的变化。', 'confidence': 0.988453209400177}
    # ]
    data =[
        {'text_region': [[8.0, 12.0], [165.0, 12.0], [165.0, 27.0], [8.0, 27.0]], 'text': '湖北省中医院', 'confidence': 0.9071481823921204},
        {'text_region': [[310.0, 14.0], [338.0, 14.0], [338.0, 23.0], [310.0, 23.0]], 'text': '北', 'confidence': 0.6512491106987}, 
        {'text_region': [[333.0, 11.0], [596.0, 11.0], [596.0, 26.0], [333.0, 26.0]], 'text': '发表论文情祝统计分机', 'confidence': 0.8921557664871216}
    ]

    res = merge_text_in_line(data)
    print(res)

    